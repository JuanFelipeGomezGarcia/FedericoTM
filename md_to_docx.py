import argparse
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_styled_paragraph(doc, text):
    p = doc.add_paragraph()
    apply_inline_styles(p, text)
    return p


def apply_inline_styles(paragraph, text):
    # Handles **bold** and *italic* (simple, non-nested)
    parts = re.split(r'(\*\*[^\*]+\*\*|\*[^\*]+\*)', text)
    for part in parts:
        if not part:
            continue
        m_b = re.match(r'^\*\*(.+)\*\*$', part)
        m_i = re.match(r'^\*(.+)\*$', part)
        if m_b:
            run = paragraph.add_run(m_b.group(1))
            run.bold = True
        elif m_i:
            run = paragraph.add_run(m_i.group(1))
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table(lines):
    # lines: list of table lines starting and ending with '|'
    rows = []
    for l in lines:
        # split on '|' but ignore leading/trailing empty from edges
        cells = [c.strip() for c in l.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        raw_lines = f.read().splitlines()

    doc = Document()
    # set default font size for normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    n = len(raw_lines)
    in_code = False
    code_lines = []

    while i < n:
        line = raw_lines[i]

        if line.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.style = 'Intense Quote'
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.rstrip()

        # Horizontal rule
        if re.match(r'^[\-\*_]{3,}$', stripped):
            doc.add_page_break()
            i += 1
            continue

        # Headings
        if stripped.lstrip().startswith('#'):
            hashes = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            level = min(hashes, 4)
            # map to Word heading levels
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Blockquote
        if stripped.lstrip().startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            p.paragraph_format.left_indent = Pt(12)
            i += 1
            continue

        # Table detection: line starts with '|' and next line is a separator with dashes
        if stripped.strip().startswith('|') and i + 1 < n and re.match(r'^\s*\|?\s*[:-]+', raw_lines[i+1]):
            table_lines = [raw_lines[i]]
            i += 1
            # consume separator line
            if i < n and re.match(r'^\s*\|?\s*[:-]+', raw_lines[i]):
                i += 1
            # following lines that start with '|' are table rows
            while i < n and raw_lines[i].strip().startswith('|'):
                table_lines.append(raw_lines[i])
                i += 1

            rows = parse_table(table_lines)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        # simple inline styles inside table
                        cell_para = cell.paragraphs[0]
                        apply_inline_styles(cell_para, cell_text)
                        if r_idx == 0:
                            for run in cell_para.runs:
                                run.bold = True
            continue

        # Unordered list
        m_list = re.match(r'^(\s*)([-\*\+]\s+)(.+)$', stripped)
        if m_list:
            text = m_list.group(3)
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_styles(p, text)
            i += 1
            continue

        # Image placeholder markers like [📸 PANTALLAZO: ...]
        m_img = re.match(r'^\[📸\s*PANTALLAZO:\s*(.+)\]$', stripped)
        if m_img:
            caption = m_img.group(1).strip()
            # insert a centered placeholder paragraph
            p = doc.add_paragraph()
            run = p.add_run('[IMAGEN: {}]'.format(caption))
            run.italic = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # add a bordered empty table as visual placeholder for image
            ph = doc.add_table(rows=1, cols=1)
            ph.style = 'Table Grid'
            ph_cell = ph.cell(0, 0)
            ph_cell.text = ''
            ph_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Empty line
        if not stripped.strip():
            doc.add_paragraph('')
            i += 1
            continue

        # Default paragraph
        add_styled_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX (improved)')
    parser.add_argument('input_md', help='Path to input Markdown file')
    parser.add_argument('output_docx', help='Path to output DOCX file')
    args = parser.parse_args()
    convert(args.input_md, args.output_docx)
